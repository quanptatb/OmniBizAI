"""
Script chuyển đổi báo cáo tốt nghiệp Markdown → Word (.docx)
Bố cục theo mẫu FPT Polytechnic DATN SP26
- Trang bìa chính thức
- Font Times New Roman 13pt
- Heading đánh số, page break giữa chương
- Sơ đồ Mermaid render thành hình ảnh
- Header/Footer + đánh số trang

Cách chạy:
  pip install python-docx requests
  python generate_docx.py
"""

import re, os, base64, requests
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(SCRIPT_DIR, "BaoCao_TotNghiep_OmniBizAI.docx")
IMG_DIR = os.path.join(SCRIPT_DIR, "mermaid_images")
os.makedirs(IMG_DIR, exist_ok=True)

FONT = "Times New Roman"
SZ = Pt(13)
SZ_SMALL = Pt(12)

# ── helpers ──────────────────────────────────────────────────────

def sf(run, size=SZ, bold=False, italic=False, color=None):
    """Set font on a run"""
    run.font.name = FONT
    run.font.size = size
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def add_p(doc, text, size=SZ, bold=False, align=None, space_after=6, italic=False):
    p = doc.add_paragraph()
    if align: p.alignment = align
    run = p.add_run(text)
    sf(run, size, bold, italic)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(22)
    return p

def add_heading_fmt(doc, text, level=1):
    sizes = {1: Pt(18), 2: Pt(15), 3: Pt(14), 4: Pt(13)}
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    sf(run, sizes.get(level, SZ), bold=True)
    # Set heading font explicitly via XML (override theme)
    rpr = run._element.get_or_add_rPr()
    rfont = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT}" w:hAnsi="{FONT}" w:eastAsia="{FONT}" w:cs="{FONT}"/>')
    rpr.append(rfont)
    return h

def add_table_grid(doc, headers, rows):
    ncols = len(headers)
    table = doc.add_table(rows=1+len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Shade header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        sf(run, SZ_SMALL, bold=True)
        # light blue shading
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci in range(min(len(row), ncols)):
            cell = table.rows[ri+1].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(row[ci]))
            sf(run, SZ_SMALL)
    doc.add_paragraph()

def render_mermaid(code, idx):
    img_path = os.path.join(IMG_DIR, f"diagram_{idx}.png")
    if os.path.exists(img_path) and os.path.getsize(img_path) > 500:
        return img_path
    try:
        encoded = base64.urlsafe_b64encode(code.encode('utf-8')).decode('ascii')
        url = f"https://mermaid.ink/img/{encoded}"
        resp = requests.get(url, timeout=30)
        if resp.status_code == 200 and len(resp.content) > 500:
            with open(img_path, 'wb') as f:
                f.write(resp.content)
            print(f"  ✅ Diagram {idx}")
            return img_path
        print(f"  ⚠️ Diagram {idx}: HTTP {resp.status_code}")
    except Exception as e:
        print(f"  ⚠️ Diagram {idx}: {e}")
    return None

def add_page_number(section):
    """Add page number to footer"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Page number field
    run = p.add_run()
    fld = parse_xml(
        f'<w:fldSimple {nsdecls("w")} w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
    )
    run._element.addnext(fld)
    sf(run, Pt(11))

# ── cover page ───────────────────────────────────────────────────

def build_cover(doc):
    """Trang bìa theo mẫu FPT Polytechnic"""
    for _ in range(2):
        add_p(doc, "", size=Pt(10))

    add_p(doc, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", Pt(14), bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "TRƯỜNG CAO ĐẲNG FPT POLYTECHNIC", Pt(16), bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(3):
        add_p(doc, "", size=Pt(10))

    add_p(doc, "BÁO CÁO", Pt(20), bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "DỰ ÁN TỐT NGHIỆP", Pt(20), bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(2):
        add_p(doc, "", size=Pt(10))

    add_p(doc, "XÂY DỰNG HỆ THỐNG QUẢN LÝ VẬN HÀNH", Pt(16), bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "THÔNG MINH CHO DOANH NGHIỆP VỪA VÀ NHỎ", Pt(16), bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "OMNIBIZAI", Pt(18), bold=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(4):
        add_p(doc, "", size=Pt(10))

    info_lines = [
        "Ngành: Ứng dụng phần mềm",
        "Giảng viên hướng dẫn: [TÊN_GVHD]",
        "",
        "Sinh viên thực hiện:",
        "    1. [Thành viên 1] - [MSSV]",
        "    2. [Thành viên 2] - [MSSV]",
        "    3. [Thành viên 3] - [MSSV]",
        "    4. [Thành viên 4] - [MSSV]",
    ]
    for line in info_lines:
        add_p(doc, line, Pt(13), align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(3):
        add_p(doc, "", size=Pt(10))

    add_p(doc, "Hà Nội, tháng 05 năm 2026", Pt(13), italic=True,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

# ── MD parser ────────────────────────────────────────────────────

def clean_md(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text

def parse_md(doc, md_text, include_mermaid=True):
    lines = md_text.split('\n')
    i = 0
    m_idx = 0
    in_table = False
    t_headers = []
    t_rows = []
    in_code = False
    code_lang = ""
    code_buf = []

    while i < len(lines):
        line = lines[i]

        # ── code block ──
        if line.strip().startswith('```') and not in_code:
            lang = line.strip().replace('```', '').strip()
            if lang:
                code_lang = lang
                code_buf = []
                in_code = True
                i += 1; continue
            i += 1; continue
        if line.strip() == '```' and in_code:
            if code_lang == 'mermaid' and include_mermaid:
                m_idx += 1
                img = render_mermaid('\n'.join(code_buf), m_idx)
                if img:
                    doc.add_picture(img, width=Inches(5.2))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_p(doc, f"Hình {m_idx}: Sơ đồ Mermaid", SZ_SMALL,
                          italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    add_p(doc, f"[Sơ đồ #{m_idx} — xem F_SoDo_Mermaid.md]", SZ_SMALL, italic=True)
            elif code_lang != 'mermaid':
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            in_code = False; code_lang = ""; code_buf = []
            i += 1; continue
        if in_code:
            code_buf.append(line)
            i += 1; continue

        # ── flush table ──
        if in_table and not line.strip().startswith('|'):
            add_table_grid(doc, t_headers, t_rows)
            in_table = False; t_headers = []; t_rows = []

        # ── table ──
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if not in_table:
                t_headers = cells; in_table = True
            elif all(set(c) <= set('-: ') for c in cells):
                pass
            else:
                t_rows.append(cells)
            i += 1; continue

        # ── headings ──
        m_h = re.match(r'^(#{1,4})\s+(.+)', line)
        if m_h:
            level = len(m_h.group(1))
            text = m_h.group(2)
            # Page break before chapter (# CHƯƠNG)
            if level == 1 and 'CHƯƠNG' in text.upper():
                doc.add_page_break()
            if level == 1 and ('PHỤ LỤC' in text.upper()):
                doc.add_page_break()
            add_heading_fmt(doc, text, min(level, 4))
            i += 1; continue

        # ── hr ──
        if line.strip() in ('---', '***', '___'):
            i += 1; continue

        # ── blockquote (skip alerts) ──
        if line.strip().startswith('>'):
            text = line.strip().lstrip('>').strip()
            if text.startswith('[!') or not text:
                i += 1; continue
            p = add_p(doc, clean_md(text), italic=True)
            p.paragraph_format.left_indent = Cm(1)
            i += 1; continue

        # ── bullet ──
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = clean_md(line.strip()[2:])
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(text)
            sf(run)
            i += 1; continue

        # ── numbered ──
        m_num = re.match(r'^(\d+)\.\s+(.+)', line.strip())
        if m_num:
            text = clean_md(m_num.group(2))
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(text)
            sf(run)
            i += 1; continue

        # ── empty ──
        if not line.strip():
            i += 1; continue

        # ── skip P2 reference ──
        if 'A_BaoCao_TotNghiep_P2' in line:
            i += 1; continue

        # ── paragraph ──
        text = clean_md(line.strip())
        if text:
            add_p(doc, text)
        i += 1

    if in_table:
        add_table_grid(doc, t_headers, t_rows)

# ── main ─────────────────────────────────────────────────────────

def main():
    print("=" * 55)
    print("📄 Tạo Word báo cáo tốt nghiệp OmniBizAI")
    print("   Font: Times New Roman 13pt")
    print("   Bố cục: FPT Polytechnic DATN SP26")
    print("=" * 55)

    doc = Document()

    # Default style
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = SZ

    # Margins (trái 3cm, phải 2cm, trên/dưới 2.5cm)
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(2)

    # Page numbers
    for sec in doc.sections:
        add_page_number(sec)

    # ── 1. TRANG BÌA ──
    print("\n📝 Trang bìa...")
    build_cover(doc)

    # ── 2. NỘI DUNG BÁO CÁO ──
    report = os.path.join(SCRIPT_DIR, "A_BaoCao_TotNghiep.md")
    if os.path.exists(report):
        print("📖 Nội dung báo cáo (Chương 1–7 + Phụ lục)...")
        with open(report, 'r', encoding='utf-8') as f:
            md = f.read()
        # Skip cover info in MD (first few lines until first ---)
        # Actually parse everything, cover is separate
        parse_md(doc, md, include_mermaid=True)
    else:
        print(f"⚠️ Không tìm thấy: {report}")

    # ── 3. PHỤ LỤC SƠ ĐỒ ──
    diagrams = os.path.join(SCRIPT_DIR, "F_SoDo_Mermaid.md")
    if os.path.exists(diagrams):
        print("📖 Phụ lục sơ đồ Mermaid...")
        doc.add_page_break()
        add_heading_fmt(doc, "PHỤ LỤC E: TỔNG HỢP SƠ ĐỒ HỆ THỐNG", 1)
        with open(diagrams, 'r', encoding='utf-8') as f:
            md = f.read()
        parse_md(doc, md, include_mermaid=True)

    # Save
    doc.save(OUTPUT)
    size_kb = os.path.getsize(OUTPUT) // 1024
    print(f"\n🎉 Hoàn tất! ({size_kb} KB)")
    print(f"📁 File: {OUTPUT}")

if __name__ == '__main__':
    main()
